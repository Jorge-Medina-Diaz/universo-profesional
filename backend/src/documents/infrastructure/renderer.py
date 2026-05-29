"""PDF + DOCX renderer.

PDF: Jinja2 → HTML → WeasyPrint.
DOCX: python-docx, programmatic.
Storage: local filesystem under STORAGE_ROOT.

ATS-friendly: simple typography, no columns, no images, normalized whitespace,
all section headers as actual <h2> tags so parsers recognize structure.
"""
from __future__ import annotations

import asyncio
import io
from pathlib import Path
from typing import Any
from uuid import UUID

import structlog
from jinja2 import Environment, FileSystemLoader, select_autoescape

from src.documents.application.ports import Renderer
from src.shared.storage import get_storage

logger = structlog.get_logger(__name__)

_TEMPLATES_DIR = Path(__file__).resolve().parents[3] / "templates" / "cv"


class WeasyPrintRenderer(Renderer):
    def __init__(self) -> None:
        self._env = Environment(
            loader=FileSystemLoader(str(_TEMPLATES_DIR)),
            autoescape=select_autoescape(["html", "j2"]),
            trim_blocks=True,
            lstrip_blocks=True,
        )

    async def render_pdf(
        self,
        *,
        content_json: dict[str, Any],
        template: str,
        language: str,
        user_id: UUID,
    ) -> str:
        try:
            tmpl_path = f"{template}.html.j2"
            tmpl = self._env.get_template(tmpl_path)
        except Exception:
            tmpl = self._env.get_template("ats-classic.html.j2")

        html = tmpl.render(resume=content_json, language=language)

        key = f"{user_id}/{_uuid4_short()}.pdf"
        try:
            from weasyprint import HTML

            # write_pdf(target=None) returns the PDF as bytes — render in a
            # thread, then hand the bytes to the storage backend (filesystem or
            # S3) so durability is the storage layer's concern, not WeasyPrint's.
            pdf_bytes: bytes = await asyncio.to_thread(
                lambda: HTML(string=html, base_url=str(_TEMPLATES_DIR)).write_pdf()
            )
            await get_storage().save(key, pdf_bytes, content_type="application/pdf")
        except Exception as exc:
            logger.warning("pdf_render_failed_fallback_html", error=str(exc))
            # Fallback: persist HTML so the document is still retrievable.
            key = key[:-4] + ".html"
            await get_storage().save(key, html.encode("utf-8"), content_type="text/html")
        return key

    async def render_docx(
        self,
        *,
        content_json: dict[str, Any],
        template: str,
        language: str,
        user_id: UUID,
    ) -> str:
        from docx import Document as DocxDocument
        from docx.shared import Pt

        doc = DocxDocument()
        styles = doc.styles
        if "Normal" in styles:
            styles["Normal"].font.name = "Calibri"
            styles["Normal"].font.size = Pt(11)

        basics = content_json.get("basics", {}) or {}

        # Cover letter: just header + body paragraphs. Skip CV sections.
        is_cover_letter = (
            template.startswith("cover-letter")
            or (content_json.get("meta", {}) or {}).get("kind") == "cover_letter"
        )
        if is_cover_letter:
            if basics.get("name"):
                doc.add_heading(basics["name"], level=1)
            if basics.get("label"):
                doc.add_paragraph(basics["label"])
            if basics.get("email"):
                doc.add_paragraph(basics["email"])
            meta = content_json.get("meta", {}) or {}
            if meta.get("target_company") or meta.get("target_title"):
                line = "Para: " + " · ".join(
                    filter(None, [meta.get("target_company"), meta.get("target_title")])
                )
                doc.add_paragraph(line)
            body = content_json.get("cover_letter_body") or basics.get("summary") or ""
            for para in body.split("\n\n"):
                doc.add_paragraph(para.strip())
            return await _save_docx(doc, user_id)

        if basics.get("name"):
            doc.add_heading(basics["name"], level=1)
        if basics.get("label"):
            doc.add_paragraph(basics["label"])
        if basics.get("summary"):
            doc.add_paragraph(basics["summary"])
        if basics.get("email"):
            doc.add_paragraph(basics["email"])

        # Experience
        if content_json.get("work"):
            doc.add_heading("Experience" if language == "en" else "Experiencia", level=2)
            for w in content_json["work"]:
                p = doc.add_paragraph()
                run = p.add_run(f"{w.get('position') or ''} — {w.get('name') or ''}")
                run.bold = True
                dates = " · ".join(filter(None, [w.get("startDate"), w.get("endDate") or "—"]))
                if dates:
                    p.add_run(f"  ({dates})")
                if w.get("summary"):
                    doc.add_paragraph(w["summary"])
                for h in w.get("highlights", []) or []:
                    doc.add_paragraph(h, style="List Bullet")

        # Education
        if content_json.get("education"):
            doc.add_heading("Education" if language == "en" else "Educación", level=2)
            for e in content_json["education"]:
                p = doc.add_paragraph()
                run = p.add_run(
                    f"{e.get('studyType') or ''} {('— ' + e['area']) if e.get('area') else ''} — {e.get('institution') or ''}"
                )
                run.bold = True
                dates = " · ".join(filter(None, [e.get("startDate"), e.get("endDate") or "—"]))
                if dates:
                    p.add_run(f"  ({dates})")

        # Skills
        if content_json.get("skills"):
            doc.add_heading("Skills" if language == "en" else "Competencias", level=2)
            doc.add_paragraph(", ".join(s.get("name", "") for s in content_json["skills"]))

        # Projects
        if content_json.get("projects"):
            doc.add_heading("Projects" if language == "en" else "Proyectos", level=2)
            for p_ in content_json["projects"]:
                p = doc.add_paragraph()
                run = p.add_run(p_.get("name") or "")
                run.bold = True
                if p_.get("description"):
                    p.add_run(" — " + p_["description"])

        # Languages
        if content_json.get("languages"):
            doc.add_heading("Languages" if language == "en" else "Idiomas", level=2)
            doc.add_paragraph(
                ", ".join(
                    f"{lang.get('language')} ({lang.get('fluency')})"
                    for lang in content_json["languages"]
                )
            )

        return await _save_docx(doc, user_id)


_DOCX_MEDIA = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


async def _save_docx(doc: Any, user_id: UUID) -> str:
    """Serialize a python-docx Document to bytes and store it; return the key."""
    buffer = io.BytesIO()
    await asyncio.to_thread(doc.save, buffer)
    key = f"{user_id}/{_uuid4_short()}.docx"
    await get_storage().save(key, buffer.getvalue(), content_type=_DOCX_MEDIA)
    return key


def _uuid4_short() -> str:
    from uuid import uuid4

    return uuid4().hex[:16]
